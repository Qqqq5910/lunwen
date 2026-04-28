from pathlib import Path
from docx import Document


def load_document(file_path):
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"找不到文件: {file_path}")
    return Document(str(path))


def _table_row_text(row):
    cells = []
    for cell in row.cells:
        text = " ".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
        if text:
            cells.append(text)
    return " ".join(cells).strip()


def _first_row_paragraph(row):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if paragraph.text.strip():
                return paragraph
    return None


def read_docx_paragraphs(document, include_tables=False):
    paragraphs = []
    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            paragraphs.append({
                "index": idx,
                "text": text,
                "paragraph": paragraph
            })
    if include_tables:
        next_index = len(paragraphs)
        for table in document.tables:
            for row in table.rows:
                text = _table_row_text(row)
                if not text:
                    continue
                paragraphs.append({
                    "index": next_index,
                    "text": text,
                    "paragraph": _first_row_paragraph(row)
                })
                next_index += 1
    return paragraphs
